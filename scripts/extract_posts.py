#!/usr/bin/env python3
"""
Maharat News Pipeline – DOCX Post Extractor
============================================
Reads every .docx in input/, splits the document into news highlights by
heading style, extracts embedded images, then writes:

  output/posts/<slug>.md          – Markdown + YAML front matter
  output/images/<hash>.<ext>      – Extracted images
  output/manifests/<doc>_manifest.json
  review/<doc>_review.csv         – Fields for manual review/enrichment

Usage:
    python scripts/extract_posts.py
    python scripts/extract_posts.py --input "input/MyFile.docx"
"""

import argparse
import csv
import hashlib
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn

# ── Paths ──────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
INPUT_DIR = ROOT / "input"
POSTS_DIR = ROOT / "data" / "posts"
IMAGES_DIR = ROOT / "data" / "images"
MANIFESTS_DIR = ROOT / "data" / "manifests"
REVIEW_DIR = ROOT / "review"


# ── Title cleaner ──────────────────────────────────────────────────────────
_TITLE_SUFFIX_RE = re.compile(
    r"\s*[\(\[]*(?:NHTI|JTC|RTC|NTC)?[\)\]]*\s*[–—\-]*\s*Internal\s*$",
    re.IGNORECASE,
)

def clean_title(title: str) -> str:
    """Strip location-code / Internal suffixes from document headings."""
    return _TITLE_SUFFIX_RE.sub("", title).strip()


# ── Slug helper (no external dep) ─────────────────────────────────────────
def slugify(text: str, max_length: int = 80) -> str:
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    text = re.sub(r"-+", "-", text).strip("-")
    return text[:max_length]


def guess_quarter(month: int) -> str:
    return f"Q{(month - 1) // 3 + 1}"


# ── Heading detection ──────────────────────────────────────────────────────
def para_heading_level(para) -> int:
    """Return heading level (1–6), or 0 if not a heading."""
    style = para.style.name if para.style else ""
    m = re.match(r"Heading\s+(\d+)", style, re.IGNORECASE)
    if m:
        return int(m.group(1))
    # Fallback: short paragraph where every non-empty run is bold
    text = para.text.strip()
    runs = [r for r in para.runs if r.text.strip()]
    if text and len(text) < 120 and runs and all(r.bold for r in runs):
        return 2
    return 0


# ── Image extraction ───────────────────────────────────────────────────────
def extract_all_images(doc, images_dir: Path) -> dict:
    """
    Pull every image from the document's relationships.
    Returns {rId: Path} for images saved to images_dir.
    """
    images_dir.mkdir(parents=True, exist_ok=True)
    saved = {}
    for rId, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue
        if rel.is_external:
            continue
        part = rel.target_part
        ext = part.content_type.split("/")[-1].replace("jpeg", "jpg")
        digest = hashlib.md5(part.blob).hexdigest()[:12]
        fname = f"{digest}.{ext}"
        out = images_dir / fname
        if not out.exists():
            out.write_bytes(part.blob)
        saved[rId] = out
    return saved


def para_image_rids(para) -> list:
    """Collect all image rIds embedded in a paragraph element."""
    ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    rids = []
    for elem in para._element.iter():
        if elem.tag == qn("a:blip"):
            rId = elem.get(f"{{{ns}}}embed")
            if rId:
                rids.append(rId)
    return rids


# ── Document segmentation ──────────────────────────────────────────────────
def segment_document(doc, split_level: int = 2) -> list:
    """
    Split document paragraphs into sections at headings <= split_level.
    Returns list of dicts: {title, level, paragraphs}.

    Tries Heading 1 first; if none found, falls back to Heading 2.
    """
    sections = []
    current = None

    for para in doc.paragraphs:
        level = para_heading_level(para)
        if level and level <= split_level:
            if current:
                sections.append(current)
            current = {"title": para.text.strip(), "level": level, "paragraphs": []}
        elif current is not None:
            current["paragraphs"].append(para)

    if current:
        sections.append(current)

    # If no headings found at this level, try a looser split
    if not sections and split_level == 1:
        return segment_document(doc, split_level=2)

    return sections


# ── Date detection ─────────────────────────────────────────────────────────
_MONTHS = (
    "January|February|March|April|May|June|July|August|"
    "September|October|November|December|"
    "Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec"
)

# Priority 1 – ISO
_PAT_ISO = re.compile(r"\b(\d{4}-\d{2}-\d{2})\b")

# Priority 2 – "Month Day[range], Year"
#   handles: "January 4, 2026"  "January 4–8, 2026"  "January 11 to 15, 2026"
#            "January 25 and 26, 2026"  "January 21-25, 2026"
_PAT_MDY = re.compile(
    rf"({_MONTHS})\.?\s+"
    r"(\d{1,2})"
    r"(?:\s*[–—\-]\s*\d{1,2}|\s+to\s+\d{1,2}|\s+and\s+\d{1,2})?"
    r",\s*(\d{4})",
    re.IGNORECASE,
)

# Priority 3 – "Month Year"  e.g. "December 2025"
_PAT_MY = re.compile(rf"\b({_MONTHS})\.?\s+(\d{{4}})\b", re.IGNORECASE)


def detect_date(paragraphs) -> tuple:
    """Return (date_str, year, quarter) from the first date-like string found."""
    for para in paragraphs:
        text = para.text

        m = _PAT_ISO.search(text)
        if m:
            try:
                dt = datetime.strptime(m.group(1), "%Y-%m-%d")
                return dt.strftime("%Y-%m-%d"), dt.year, guess_quarter(dt.month)
            except ValueError:
                pass

        m = _PAT_MDY.search(text)
        if m:
            month_s, day_s, year_s = m.group(1), m.group(2), m.group(3)
            for fmt in ("%B %d %Y", "%b %d %Y"):
                try:
                    dt = datetime.strptime(f"{month_s} {day_s} {year_s}", fmt)
                    return dt.strftime("%Y-%m-%d"), dt.year, guess_quarter(dt.month)
                except ValueError:
                    pass

        m = _PAT_MY.search(text)
        if m:
            month_s, year_s = m.group(1), m.group(2)
            for fmt in ("%B %Y", "%b %Y"):
                try:
                    dt = datetime.strptime(f"{month_s} {year_s}", fmt)
                    return dt.strftime("%Y-%m-%d"), dt.year, guess_quarter(dt.month)
                except ValueError:
                    pass

    return "", 0, ""


# ── Markdown body builder ──────────────────────────────────────────────────
def paragraphs_to_markdown(paragraphs) -> str:
    lines = []
    for para in paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = para.style.name if para.style else ""
        level = para_heading_level(para)
        if level:
            lines.append(f"{'#' * level} {text}")
        elif re.match(r"List", style, re.IGNORECASE):
            lines.append(f"- {text}")
        else:
            lines.append(text)
    # Collapse 3+ blank lines → 2
    body = re.sub(r"\n{3,}", "\n\n", "\n".join(lines))
    return body.strip()


# ── Build a single post dict ───────────────────────────────────────────────
def build_post(section: dict, index: int, source_doc: str, rid_to_path: dict) -> dict:
    title = clean_title(section["title"])
    # If heading was blank, use first body paragraph as title
    if not title:
        for p in section["paragraphs"]:
            if p.text.strip():
                title = p.text.strip()[:80]
                break
    if not title:
        title = f"Untitled Post {index:03d}"
    slug = slugify(title) or f"post-{index:03d}"

    # Collect images referenced in this section's paragraphs
    image_rids = []
    for para in section["paragraphs"]:
        image_rids.extend(para_image_rids(para))

    # Deduplicate while preserving order
    seen = set()
    unique_rids = []
    for r in image_rids:
        if r not in seen:
            seen.add(r)
            unique_rids.append(r)

    image_rel_paths = [
        f"images/{rid_to_path[r].name}"
        for r in unique_rids
        if r in rid_to_path
    ]
    featured_image = image_rel_paths[0] if image_rel_paths else ""
    gallery_images = image_rel_paths[1:] if len(image_rel_paths) > 1 else []

    date_str, year, quarter = detect_date(section["paragraphs"])

    # Summary: first paragraph with real text, truncated
    summary = ""
    for para in section["paragraphs"]:
        t = para.text.strip()
        if t:
            summary = t[:200] + ("…" if len(t) > 200 else "")
            break

    body_md = paragraphs_to_markdown(
        [p for p in section["paragraphs"] if p.text.strip()]
    )

    return {
        "title": title,
        "internal": f"post-{index:03d}",
        "slug": slug,
        "date": date_str,
        "year": year,
        "quarter": quarter,
        "summary": summary,
        "body_markdown": body_md,
        "category": "",
        "tags": [],
        "location": "",
        "partner": "",
        "featured_image": featured_image,
        "gallery_images": gallery_images,
        "seo_title": title[:60],
        "seo_description": summary[:155] if summary else "",
        "source_document": source_doc,
        "source_section": title,
        "source_page": 0,
    }


# ── Writers ────────────────────────────────────────────────────────────────
def write_markdown(post: dict, posts_dir: Path) -> Path:
    front_matter = {k: v for k, v in post.items() if k != "body_markdown"}
    body = post["body_markdown"]
    content = (
        "---\n"
        + yaml.dump(front_matter, allow_unicode=True, sort_keys=False)
        + "---\n\n"
        + body
        + "\n"
    )
    out = posts_dir / f"{post['slug']}.md"
    out.write_text(content, encoding="utf-8")
    return out


def write_manifest(posts: list, manifests_dir: Path, source_doc: str) -> Path:
    manifest = {
        "source": source_doc,
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "count": len(posts),
        "posts": [
            {
                "internal": p["internal"],
                "slug": p["slug"],
                "title": p["title"],
                "date": p["date"],
                "featured_image": p["featured_image"],
                "image_count": len(p["gallery_images"]) + (1 if p["featured_image"] else 0),
            }
            for p in posts
        ],
    }
    out = manifests_dir / f"{Path(source_doc).stem}_manifest.json"
    out.write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")
    return out


REVIEW_FIELDS = [
    "internal", "slug", "title", "date", "year", "quarter",
    "category", "tags", "location", "partner",
    "seo_title", "seo_description", "featured_image", "summary",
]


def write_review_csv(posts: list, review_dir: Path, source_doc: str) -> Path:
    out = review_dir / f"{Path(source_doc).stem}_review.csv"
    with out.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=REVIEW_FIELDS, extrasaction="ignore")
        writer.writeheader()
        for p in posts:
            row = dict(p)
            row["tags"] = ", ".join(p.get("tags", []))
            writer.writerow(row)
    return out


# ── Process one DOCX ──────────────────────────────────────────────────────
def process_docx(docx_path: Path, split_level: int = 1) -> list:
    print(f"\n── {docx_path.name}")
    doc = Document(str(docx_path))

    rid_to_path = extract_all_images(doc, IMAGES_DIR)
    print(f"   Images extracted : {len(rid_to_path)}")

    sections = segment_document(doc, split_level=split_level)
    print(f"   Sections found   : {len(sections)}")

    # Drop sections with no title and no body text
    sections = [
        s for s in sections
        if s["title"].strip() or any(p.text.strip() for p in s["paragraphs"])
    ]

    if not sections:
        print("   WARNING: No headings detected — wrapping whole doc as one post.")
        sections = [{"title": docx_path.stem, "level": 1, "paragraphs": list(doc.paragraphs)}]

    POSTS_DIR.mkdir(parents=True, exist_ok=True)
    MANIFESTS_DIR.mkdir(parents=True, exist_ok=True)
    REVIEW_DIR.mkdir(parents=True, exist_ok=True)

    posts = []
    for i, section in enumerate(sections, start=1):
        post = build_post(section, i, docx_path.name, rid_to_path)
        posts.append(post)
        md_path = write_markdown(post, POSTS_DIR)
        img_count = len(post["gallery_images"]) + (1 if post["featured_image"] else 0)
        print(f"   [{i:02d}] {post['slug']}.md  ({img_count} image(s))")

    manifest_path = write_manifest(posts, MANIFESTS_DIR, docx_path.name)
    csv_path = write_review_csv(posts, REVIEW_DIR, docx_path.name)
    print(f"\n   Manifest → {manifest_path.relative_to(ROOT)}")
    print(f"   Review   → {csv_path.relative_to(ROOT)}")

    return posts


# ── CLI entry point ────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="Extract news posts from DOCX files.")
    parser.add_argument(
        "--input", "-i",
        help="Path to a specific .docx file (default: all files in input/)",
    )
    parser.add_argument(
        "--split-level", "-l",
        type=int,
        default=1,
        choices=[1, 2, 3],
        help="Heading level used to split sections (default: 1). "
             "Use 2 if Heading 1 is a document title and Heading 2 marks each story.",
    )
    args = parser.parse_args()

    if args.input:
        docx_files = [Path(args.input)]
    else:
        docx_files = sorted(INPUT_DIR.glob("*.docx"))

    if not docx_files:
        print(f"No .docx files found in {INPUT_DIR}")
        sys.exit(1)

    total = 0
    for f in docx_files:
        posts = process_docx(f, split_level=args.split_level)
        total += len(posts)

    print(f"\n✓ Done. Total posts extracted: {total}")


if __name__ == "__main__":
    main()
